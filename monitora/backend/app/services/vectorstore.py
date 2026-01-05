"""MONITORA Backend - Multi-tenant RAG Pipeline"""
import os
import time
import hashlib
from typing import List, Tuple, Optional, Dict, Any
from pathlib import Path
import chromadb
from chromadb.config import Settings as ChromaSettings
from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from sentence_transformers import SentenceTransformer

from app.core.config import settings


class EmbeddingService:
    """Shared embedding service across all workspaces."""
    
    _instance: Optional['EmbeddingService'] = None
    
    def __new__(cls):
        if cls._instance is None:
            cls._instance = super().__new__(cls)
            cls._instance._initialized = False
        return cls._instance
    
    def __init__(self):
        if self._initialized:
            return
        print(f"🔧 Loading embedding model: {settings.embedding_model}")
        self.model = SentenceTransformer(settings.embedding_model)
        self._initialized = True
        print(f"✅ Embedding model loaded")
    
    def embed_text(self, text: str) -> List[float]:
        """Generate embedding for a single text."""
        embedding = self.model.encode(text, convert_to_tensor=False)
        return embedding.tolist()
    
    def embed_texts(self, texts: List[str]) -> List[List[float]]:
        """Generate embeddings for multiple texts."""
        embeddings = self.model.encode(texts, convert_to_tensor=False, show_progress_bar=True)
        return embeddings.tolist()


class WorkspaceVectorStore:
    """Vector store for a single workspace."""
    
    def __init__(self, workspace_id: str, embedding_service: EmbeddingService):
        self.workspace_id = workspace_id
        self.embedding_service = embedding_service
        
        # Create workspace-specific directory
        self.persist_directory = os.path.join(
            settings.vectorstore_base_path, 
            workspace_id
        )
        os.makedirs(self.persist_directory, exist_ok=True)
        
        # Initialize ChromaDB client for this workspace
        self.client = chromadb.PersistentClient(
            path=self.persist_directory,
            settings=ChromaSettings(anonymized_telemetry=False)
        )
        
        # Get or create collection
        self.collection = self.client.get_or_create_collection(
            name="documents",
            metadata={"hnsw:space": "cosine"}
        )
    
    def add_documents(self, documents: List[Document]) -> int:
        """Add documents to the vector store."""
        if not documents:
            return 0
        
        texts = [doc.page_content for doc in documents]
        metadatas = [doc.metadata for doc in documents]
        
        # Generate embeddings
        embeddings = self.embedding_service.embed_texts(texts)
        
        # Generate unique IDs
        ids = [f"doc_{i}_{hashlib.md5(text[:100].encode()).hexdigest()[:8]}" 
               for i, text in enumerate(texts)]
        
        # Add to collection
        self.collection.add(
            embeddings=embeddings,
            documents=texts,
            metadatas=metadatas,
            ids=ids
        )
        
        return len(documents)
    
    def similarity_search(self, query: str, k: int = 5) -> List[Tuple[Document, float]]:
        """Search for similar documents."""
        query_embedding = self.embedding_service.embed_text(query)
        
        results = self.collection.query(
            query_embeddings=[query_embedding],
            n_results=k
        )
        
        documents_with_scores = []
        if results['documents'] and results['documents'][0]:
            for doc_text, metadata, distance in zip(
                results['documents'][0],
                results['metadatas'][0],
                results['distances'][0]
            ):
                doc = Document(page_content=doc_text, metadata=metadata)
                # Convert distance to similarity score
                score = 1 - distance
                documents_with_scores.append((doc, score))
        
        return documents_with_scores
    
    def count(self) -> int:
        """Get document count."""
        return self.collection.count()
    
    def clear(self) -> None:
        """Clear all documents."""
        self.client.delete_collection("documents")
        self.collection = self.client.create_collection(
            name="documents",
            metadata={"hnsw:space": "cosine"}
        )


class VectorStoreManager:
    """Manager for multi-tenant vector stores."""
    
    def __init__(self):
        self.embedding_service = EmbeddingService()
        self._stores: Dict[str, WorkspaceVectorStore] = {}
    
    def get_store(self, workspace_id: str) -> WorkspaceVectorStore:
        """Get or create vector store for a workspace."""
        if workspace_id not in self._stores:
            self._stores[workspace_id] = WorkspaceVectorStore(
                workspace_id=workspace_id,
                embedding_service=self.embedding_service
            )
        return self._stores[workspace_id]
    
    def delete_store(self, workspace_id: str) -> bool:
        """Delete a workspace's vector store."""
        if workspace_id in self._stores:
            del self._stores[workspace_id]
        
        # Delete directory
        store_path = os.path.join(settings.vectorstore_base_path, workspace_id)
        if os.path.exists(store_path):
            import shutil
            shutil.rmtree(store_path)
            return True
        return False


# Global manager instance
vectorstore_manager = VectorStoreManager()


class DocumentProcessor:
    """Process documents into chunks for indexing."""
    
    def __init__(self, chunk_size: int = 1500, chunk_overlap: int = 300):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", ". ", " ", ""]
        )
    
    def process_text(self, text: str, metadata: Dict[str, Any]) -> List[Document]:
        """Split text into chunks."""
        chunks = self.text_splitter.split_text(text)
        return [
            Document(page_content=chunk, metadata={**metadata, "chunk_index": i})
            for i, chunk in enumerate(chunks)
        ]
    
    def process_pdf(self, file_path: str, metadata: Dict[str, Any]) -> List[Document]:
        """Extract and chunk text from PDF."""
        import pdfplumber
        
        text = ""
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except Exception as e:
            print(f"Error processing PDF: {e}")
            return []
        
        return self.process_text(text, metadata)
    
    def process_docx(self, file_path: str, metadata: Dict[str, Any]) -> List[Document]:
        """Extract and chunk text from DOCX."""
        from docx import Document as DocxDocument
        
        try:
            doc = DocxDocument(file_path)
            text = "\n".join([para.text for para in doc.paragraphs])
        except Exception as e:
            print(f"Error processing DOCX: {e}")
            return []
        
        return self.process_text(text, metadata)
    
    def process_txt(self, file_path: str, metadata: Dict[str, Any]) -> List[Document]:
        """Process plain text file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
        except Exception as e:
            print(f"Error processing TXT: {e}")
            return []
        
        return self.process_text(text, metadata)
    
    def process_file(self, file_path: str, file_type: str, metadata: Dict[str, Any]) -> List[Document]:
        """Process a file based on its type."""
        processors = {
            'pdf': self.process_pdf,
            'docx': self.process_docx,
            'txt': self.process_txt,
            'md': self.process_txt,
        }
        
        processor = processors.get(file_type.lower())
        if not processor:
            raise ValueError(f"Unsupported file type: {file_type}")
        
        return processor(file_path, metadata)
